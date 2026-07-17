#!/usr/bin/env python3
"""Create a flat numbered PDF folder and a matching DOCX title index."""

from __future__ import annotations

import argparse
import json
import os
import re
import unicodedata
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOC_FONT = "Times New Roman"


VENUES = [
    "AAAI", "CVPR", "ICCV", "ECCV", "NeurIPS", "ICLR", "ICML",
    "TPAMI", "TIP", "TNNLS", "TMM", "TCSVT", "Other IEEE Transactions",
]
VENUE_ORDER = {name: index for index, name in enumerate(VENUES)}

CONTENT_TAG_RULES = [
    ("Medical Imaging", r"medical|radiograph|echocardi|\bmri\b|chest"),
    ("Autonomous Driving", r"autonomous driv|self-driv|driving world|driving scene|end-to-end driving|learning to drive|dream to drive|driveworld|drivedreamer|drivingdojo|raw2drive|world4drive|\blidar\b|\bbev\b|occupanc\w*|traffic|\bcarla\b|autonomous vehicle|vehicle dynamics"),
    ("Robot Manipulation", r"robot\w*|manipulat\w*|dexter\w*|humanoid|visuomotor|vision-language-action|\bvlas?\b"),
    ("Embodied Navigation", r"navigation|waypoint|vision-language navigation|\bvln\b|goal navigation"),
    ("Video Modeling / Generation", r"video world model|video modeling|video generat\w*|generative video|video synthesis|world generation|world simulator|visual forecasting|future frame"),
    ("3D / 4D Modeling", r"\b3d\b|\b4d\b|gaussian|voxel|ray space|scene reconstruction|scene modeling|point cloud"),
    ("Reinforcement Learning", r"reinforcement learning|\brl\b|\bpolicy\b|continuous control|\batari\b"),
    ("Planning / Exploration", r"planning|exploration|imagination|trajectory optimization|goal-conditioned"),
    ("Representation Learning", r"representation|tokeniz\w*|token-based|latent|codebook|masked world model|disentangl\w*"),
    ("Object-Centric / Causal", r"object-centric|causal|causality|entit\w*|factoriz\w*|compositional"),
    ("Physical Dynamics", r"physic\w*|dynamics|simulation|simulator|motion|physical interaction|human-scene interaction|motor adaptation"),
    ("Language / Multimodal", r"language model|large language|\bllm\b|\bvlm\b|multimodal|vision-language"),
    ("Multi-Agent Systems", r"multi-agent|multi-robot|cooperation|collaborative"),
    ("Games / Interactive Environments", r"\batari\b|\bgame\w*\b|interactive environment|playerone"),
    ("Human / Social Dynamics", r"human motion|human-scene|social dynamics|social world|virtual human"),
    ("Safety / Robustness", r"safe\w*|robust\w*|generaliz\w*|uncertain\w*|distract\w*|out-of-distribution"),
    ("Survey / Benchmark", r"survey|review|benchmark|evaluat\w*|analysis"),
    ("Remote Sensing", r"remote sensing|earth observation|global-scale"),
    ("Wireless Networks", r"wireless|channel map|spectrum sharing|\b6g\b|uav rescue network"),
]


def infer_content_tags(title: str, limit: int = 3) -> str:
    normalized = title.lower()
    tags = [label for label, pattern in CONTENT_TAG_RULES if re.search(pattern, normalized)]
    if not tags:
        tags = ["World Model Learning"]
    return " · ".join(tags[:limit])


def safe_name(value: str, limit: int = 145) -> str:
    value = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode()
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("_")
    return value[:limit].rstrip("_.") or "paper"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    width.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_row_cant_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    no_split = OxmlElement("w:cantSplit")
    properties.append(no_split)


def set_fixed_table_grid(table, widths_cm: list[float]) -> None:
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(int(Cm(sum(widths_cm)).twips)))
    table_width.set(qn("w:type"), "dxa")

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_cm in widths_cm:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(int(Cm(width_cm).twips)))
        grid.append(column)


def set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, east_asia: str, latin: str, size: float, bold: bool = False, color: str = "222222") -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_run_font(run, DOC_FONT, DOC_FONT, 9, color="666666")


def make_docx(rows: list[dict], output: Path, downloaded: int) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = DOC_FONT
    normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("World Model Paper Title Index (2023-2026)")
    set_run_font(run, DOC_FONT, DOC_FONT, 18, bold=True, color="17365D")

    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_after = Pt(10)
    if downloaded == len(rows):
        summary_text = f"{len(rows)} papers; standardized content tags inferred from paper titles"
    else:
        summary_text = (
            f"{len(rows)} papers; {downloaded} PDFs downloaded and validated; "
            f"{len(rows) - downloaded} open-access PDFs unavailable"
        )
    run = summary.add_run(summary_text)
    set_run_font(run, DOC_FONT, DOC_FONT, 9.5, color="555555")

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    widths = [1.00, 1.20, 2.20, 8.80, 4.80]
    set_fixed_table_grid(table, widths)
    headers = ["No.", "Year", "Venue", "Paper title", "Content tags"]
    header = table.rows[0]
    set_repeat_header(header)
    for index, (cell, text) in enumerate(zip(header.cells, headers)):
        set_cell_width(cell, widths[index])
        set_cell_margins(cell)
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, DOC_FONT, DOC_FONT, 9, bold=True, color="17365D")

    for row_index, item in enumerate(rows, 1):
        row = table.add_row()
        set_row_cant_split(row)
        cells = row.cells
        values = [
            item["display_no"],
            str(item["year"]),
            item["venue"],
            item["title"],
            item.get("content_tags") or infer_content_tags(item["title"]),
        ]
        for index, (cell, text) in enumerate(zip(cells, values)):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                set_cell_shading(cell, "F7FAFC")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in (3, 4) else WD_ALIGN_PARAGRAPH.CENTER
            size = 8.8 if index == 3 else (8.0 if index == 4 else 8.5)
            if index == 4:
                tags = text.split(" · ")
                for tag_index, tag in enumerate(tags):
                    if tag_index:
                        separator = paragraph.add_run(" · ")
                        set_run_font(separator, DOC_FONT, DOC_FONT, size, color="222222")
                    run = paragraph.add_run(tag)
                    color = "C00000" if tag == "Video Modeling / Generation" else "222222"
                    set_run_font(run, DOC_FONT, DOC_FONT, size, color=color)
            else:
                run = paragraph.add_run(text)
                set_run_font(run, DOC_FONT, DOC_FONT, size, color="222222")

    add_page_number(section.footer.paragraphs[0])
    properties = document.core_properties
    properties.title = "世界模型文献题名索引（2023—2026）"
    properties.subject = "AAAI、CVPR、ICCV、ECCV、NeurIPS、ICLR、ICML及IEEE Transactions世界模型文献"
    properties.author = "DoR Literature Audit"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", required=True)
    parser.add_argument("--download_status", required=True)
    parser.add_argument("--flat_dir", required=True)
    parser.add_argument("--docx", required=True)
    parser.add_argument("--index_json", required=True)
    args = parser.parse_args()
    manifest_path = Path(args.manifest)
    payload = json.loads(manifest_path.read_text())
    status = {row["id"]: row for row in json.loads(Path(args.download_status).read_text())["papers"]}
    all_rows = sorted(
        payload["papers"],
        key=lambda row: (int(row["year"]), VENUE_ORDER.get(row["venue"], 999), row["title"].lower()),
    )
    rows = []
    for original in all_rows:
        item = status.get(original["id"], {})
        source = Path(item.get("path", "")) if "downloaded" in item.get("status", "") else None
        if not source or not source.exists():
            continue
        row = dict(original)
        row["source_pdf_path"] = str(source)
        row["download_status"] = item.get("status", "downloaded")
        rows.append(row)
    flat_dir = Path(args.flat_dir)
    flat_dir.mkdir(parents=True, exist_ok=True)
    expected = set()
    for number, row in enumerate(rows, 1):
        display_no = f"{number:03d}"
        source = Path(row.pop("source_pdf_path"))
        target = flat_dir / f"{display_no}_{row['year']}_{safe_name(row['venue'], 30)}_{safe_name(row['title'])}.pdf"
        expected.add(target.name)
        if not target.exists():
            os.link(source, target)
        row["display_no"] = display_no
        row["flat_pdf_path"] = str(target)
    stale = [path for path in flat_dir.glob("*.pdf") if path.name not in expected]
    if stale:
        raise RuntimeError(f"stale numbered PDFs detected: {len(stale)}; clean manually before rerun")
    index_payload = {
        "count": len(rows),
        "downloaded": sum(bool(row["flat_pdf_path"]) for row in rows),
        "missing": sum(not row["flat_pdf_path"] for row in rows),
        "papers": [{key: row.get(key, "") for key in ("display_no", "year", "venue", "title", "id", "flat_pdf_path", "download_status")} for row in rows],
    }
    Path(args.index_json).write_text(json.dumps(index_payload, ensure_ascii=False, indent=2) + "\n")
    tsv = Path(args.index_json).with_suffix(".tsv")
    tsv.write_text("编号\t年份\t来源\t文献题名\t全文状态\tPDF文件\n" + "".join(
        f"{row['display_no']}\t{row['year']}\t{row['venue']}\t{row['title'].replace(chr(9), ' ')}\t{'已下载' if row['flat_pdf_path'] else '未获得开放全文'}\t{Path(row['flat_pdf_path']).name if row['flat_pdf_path'] else ''}\n" for row in rows
    ))
    make_docx(rows, Path(args.docx), index_payload["downloaded"])
    print(f"[done] indexed={len(rows)} linked={index_payload['downloaded']} missing={index_payload['missing']} docx={args.docx}")


if __name__ == "__main__":
    main()
